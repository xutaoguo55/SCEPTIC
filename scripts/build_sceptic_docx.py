"""
SCEPTIC framework paper for Scientific Reports.
v3 — clean layout: figures flow inline, tables after paragraphs, section-level page breaks only.
"""
import os
import csv
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTDIR = '/Users/guoxutao/sc_analysis/figures_crm'
BASE  = '/Users/guoxutao/sc_analysis'

doc = Document()

# --- styles ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
# reduce default paragraph spacing
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(0)

for sec in doc.sections:
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)

def h(text, level=1):
    hd = doc.add_heading(text, level=level)
    hd.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hd.paragraph_format.space_before = Pt(10)
    hd.paragraph_format.space_after = Pt(4)
    for r in hd.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

def para(text, bold=False, italic=False, size=11):
    """body paragraph"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p

def page_break():
    doc.add_page_break()

def fmt_cell(val):
    """Compact numeric display for manuscript tables; CSV files retain precision."""
    text = str(val)
    try:
        num = float(text)
    except ValueError:
        return text
    if num == 0:
        return '0'
    if abs(num) < 0.001:
        return f'{num:.2e}'
    if abs(num) >= 100:
        return f'{num:.0f}'
    return f'{num:.3f}'.rstrip('0').rstrip('.')

def insert_figure(path, w=6.0, caption=None):
    """insert a figure; if caption supplied, label it as a Figure paragraph"""
    if not os.path.exists(path):
        print(f"  WARNING: missing {path}")
        return
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_before = Pt(8)
    par.paragraph_format.space_after = Pt(2)
    par.add_run().add_picture(path, width=Inches(w))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(8)
        r = cp.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)

def add_table(headers, rows, caption=None, col_widths=None):
    """add a formatted table"""
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cp.paragraph_format.space_before = Pt(8)
        cp.paragraph_format.space_after = Pt(2)
        r = cp.add_run(caption)
        r.bold = True
        r.font.size = Pt(9)

    ncols = len(headers)
    nrows = len(rows) + 1
    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # header row
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        pp = cell.paragraphs[0]
        run = pp.add_run(header)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1F3A5F')
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)

    # data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            pp = cell.paragraphs[0]
            run = pp.add_run(fmt_cell(val))
            run.font.size = Pt(8)
            if i % 2 == 1:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'F0F4F8')
                shading.set(qn('w:val'), 'clear')
                cell._tc.get_or_add_tcPr().append(shading)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    # spacer after table
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(2)
    return table

def add_csv_table(path, caption, col_widths=None, max_rows=None):
    """Insert a compact CSV table from a verified supplementary data file."""
    if not os.path.exists(path):
        para(f'Missing supplementary data file: {path}', italic=True)
        return None
    with open(path, newline='', encoding='utf-8-sig') as fh:
        reader = csv.reader(fh)
        headers = next(reader)
        rows = list(reader)
    if max_rows is not None:
        rows = rows[:max_rows]
    return add_table(headers, rows, caption, col_widths=col_widths)

# ============================================================
# TITLE PAGE
# ============================================================
ti = doc.add_paragraph()
ti.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ti.add_run('SCEPTIC: a multi-layer null-control framework for validating\n'
               'single-cell transcription factor co-expression claims')
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

doc.add_paragraph()  # spacer

au = doc.add_paragraph()
au.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = au.add_run('Xiaolei Wei¹†, Haiqing Zheng²†, Hongbing Jiang³†, Junwei Huang¹,\n'
               'Qi Wei¹, Yongqiang Wei¹, Ru Feng¹, Xutao Guo¹,⁴*')
r.font.size = Pt(12)

doc.add_paragraph()

af = doc.add_paragraph()
af.alignment = WD_ALIGN_PARAGRAPH.LEFT
af.add_run('¹ Department of Hematology, Nanfang Hospital, Southern Medical University, Guangzhou 510515, China\n'
           '² Department of Nosocomial Infection Management, Nanfang Hospital, Southern Medical University, Guangzhou 510515, China\n'
           '³ Jieyang People\'s Hospital, Jieyang Cancer Center, Jieyang 522000, China\n'
           '⁴ Clinical Medical Research Center of Hematological Diseases of Guangdong Province, Guangzhou 510515, China\n\n'
           '† These authors contributed equally.\n'
           '* Correspondence: Xutao Guo (gxt827@126.com, ORCID: 0000-0001-6191-2204)').font.size = Pt(10)

# ============================================================
# ABSTRACT
# ============================================================
h('Abstract')
para('Single-cell transcriptomic studies frequently report transcription factor (TF) '
     'co-expression networks from per-sample aggregate scores¹⁻⁴, but significant '
     'TF-score correlations can arise from scoring method, global expression, cell-type '
     'composition, batch effects, or disease severity⁵⁻⁸. We present SCEPTIC '
     '(Single-Cell Expression Program Testing and Integrity Checklist), a seven-layer '
     'null-control framework that asks a simple question before a TF-score correlation '
     'is interpreted biologically: does the claim survive orthogonal scoring, independent '
     'context checks, empirical TF-null comparison, composition adjustment, cohort '
     'sensitivity, and label permutation? We stress-tested SCEPTIC on a candidate '
     'STAT4/NFKB1/GATA1 (HCTL) triad nominated from hyperinflammatory-disease PBMC '
     'single-cell data, across 1,161 patients, 6 cohorts, and 3 assay types. The '
     'framework did not support the triad as a general cross-disease regulatory axis. '
     'NFKB1–GATA1 collapsed under AUCell scoring (r = +0.15, p = 0.33), reversed in '
     'bulk blood (r = −0.07), and was weaker than random TF triplets (empirical '
     'p = 0.60). STAT4–GATA1 showed no reproducible PBMC signal and was confounded by '
     'erythroid GATA1 targets. STAT4–NFKB1 retained PBMC-context support in independent '
     'single-cell and cytokine-level datasets (SCP548 r = +0.56; E-MTAB-10026 r = +0.54; '
     'CyTOF pathway r = +0.80), but failed bulk replication and leave-one-cohort-out '
     'robustness. Thus, SCEPTIC narrowed an initially broad triad claim to one '
     'context-dependent PBMC covariation and two rejected pairs. Applying SCEPTIC L4 to '
     '15 published TF relationships showed that 7 of 15 did not exceed the 75th '
     'percentile of PBMC TF co-variation, illustrating that valid biological TF '
     'relationships do not necessarily yield pair-specific per-sample correlations in '
     'mixed tissues. We provide SCEPTIC as an open-source checklist and recommend layers '
     '1 and 4 as minimum reporting checks for TF-score co-expression claims.')

# ============================================================
# INTRODUCTION
# ============================================================
page_break()
h('Introduction')

para('Single-cell RNA sequencing (scRNA-seq) has enabled discovery of cell-type-specific '
     'transcriptional programs across diverse biological contexts¹⁻⁴. A common downstream '
     'analysis infers per-sample transcription factor (TF) activity using gene set '
     'scoring, SCENIC, DoRothEA, decoupleR, or VIPER⁵⁻⁸ and reports pairwise or '
     'network-level co-variation as evidence of coordinated activity⁹. Such analyses '
     'underpin claims from disease-specific regulatory programs to putative '
     'cross-disease signatures¹⁰.')

para('However, per-sample TF target-gene mean scoring has known vulnerabilities. '
     'Target-gene expression can be confounded by library size, cell-type composition, '
     'global transcriptional activity, and batch effects¹¹⁻¹⁵. A statistically '
     'significant Pearson correlation between two TF scores does not distinguish '
     'cell-intrinsic co-regulation from the simple observation that both pathways rise '
     'with disease severity—especially in heterogeneous PBMC samples. Despite growing '
     'awareness of reproducibility challenges in single-cell data science¹⁶, systematic '
     'null-model frameworks for validating TF co-expression claims are not widely '
     'adopted in the single-cell literature.')

para('We use "TF co-expression" throughout to refer specifically to the correlation '
     'of per-sample TF target-gene aggregate scores, and distinguish this from TF '
     'activity, which would require direct measurement (e.g., by phospho-flow, ChIP-seq, '
     'or nuclear translocation assays).')

para('Here we present SCEPTIC (Single-Cell Expression Program Testing and Integrity '
     'Checklist), a seven-layer null-control framework designed to test TF co-expression '
     'claims before they are reported as biological discoveries. The paper has three '
     'linked aims: first, to define the validation layers and their pass/fail logic; '
     'second, to stress-test the framework on a deliberately vulnerable candidate '
     'STAT4/NFKB1/GATA1 transcriptional triad (hereafter HCTL, reflecting the originating '
     'hypercytokinemia–T-lymphocyte clinical context) nominated from our in-house '
     'hyperinflammatory-disease single-cell analysis; and third, to ask whether the same '
     'empirical-null principle changes interpretation of published TF co-expression '
     'claims. The HCTL triad is therefore used as a case study for claim validation, not '
     'as a pre-assumed disease mechanism. Across 1,161 patients, 6 cohorts, and 3 assay '
     'types spanning sepsis¹⁷⁻¹⁹ and COVID-19²⁰ contexts, we use SCEPTIC to separate a '
     'PBMC-context STAT4–NFKB1 signal from GATA1-involving artifacts and then apply L4 '
     'to 15 peer-reviewed TF relationships to demonstrate its use with '
     'context-specific null designs.')

# ============================================================
# RESULTS
# ============================================================
page_break()
h('Results')

h('The SCEPTIC framework', 2)
para('SCEPTIC comprises seven validation layers, ordered from least to most stringent '
     '(Table 1, Figure 1). A TF co-expression claim is considered supported if it '
     'passes layers 1–4 and at least one of layers 5–7 in independent data; '
     'method-dependent if it passes layers 1–2 in the discovery context but fails at '
     'layers 3–4; and rejected if it fails at layer 1 (orthogonal scoring) and '
     'layer 4 (random null). These thresholds are reporting heuristics for '
     'standardised triage, not universal biological cutoffs; full parameter choices '
     'and statistical families are reported in Supplementary Tables S6-S7.')

# Table 1
t1_headers = ['Layer', 'Name', 'What it tests', 'Method', 'Pass criteria']
t1_rows = [
    ['L1', 'Orthogonal scoring', 'Method-dependence of raw mean correlation',
     'Pearson r under raw target-gene mean + regulon-based scoring (AUCell/DoRothEA)',
     'Same sign + p < 0.05 in both, |r| > 0.3'],
    ['L2', 'Cross-assay replication\n(protein-level, supportive)', 'Transcript-level finding at protein level',
     'Plasma proteomics (Olink/SomaScan) or CyTOF/flow intracellular staining',
     'Consistent direction at protein level (supportive, not confirmatory)'],
    ['L3', 'Bulk-tissue replication', 'Cross-platform / bulk-context consistency',
     'Pearson r in independent bulk RNA-seq/microarray cohort',
     'Same sign + p < 0.05, |r| > 0.15 in bulk'],
    ['L4', 'Random TF-pair null', 'Specificity vs. random TF pairs',
     'All N(N-1)/2 pairwise TF-score |r| as empirical null; empirical p = fraction with |r| ≥ observed',
     'Observed |r| > 95th percentile of null'],
    ['L5', 'Composition-adjusted r', 'Cell-type proportion confounding',
     'Partial r after regressing out cell-type proportions + housekeeping mean',
     'Partial |r| > 0.3, p < 0.05'],
    ['L6', 'Leave-one-cohort-out', 'Single-cohort dependence',
     'Random-effects meta-analysis (DerSimonian-Laird) as heterogeneity sensitivity',
     'Same sign + significance in all LOO; report I²/tau²/Q'],
    ['L7', 'Label permutation\n(supportive check)', 'Disease-group dependence of correlation',
     '10,000 disease-label permutations per cohort',
     'Observed |r| > 95th percentile of null'],
]
add_table(t1_headers, t1_rows, 'Table 1. The seven SCEPTIC validation layers.',
          col_widths=[0.35, 0.9, 1.1, 2.0, 1.3])

# Figure 1
insert_figure(os.path.join(OUTDIR, 'Fig1_final.png'), w=6.5,
              caption='Figure 1. The SCEPTIC framework. (A) Per-sample TF-score correlations can '
                      'be confounded by global expression, cell-type composition, batch/cohort '
                      'effects, and disease severity. (B) The seven SCEPTIC validation layers. '
                      '(C) HCTL triad case-study verdicts. Green = passed, orange = '
                      'supportive/context-dependent, red = not specific.')

para('The results are organised to mirror this decision path. We first define the '
     'seven-layer framework, then apply each layer to the HCTL triad, using the '
     'layer-by-layer failures to explain why an initially significant correlation is '
     'downgraded, rejected, or retained only as context-dependent evidence. We then '
     'separate the framework from the case study by applying the L4 empirical-null '
     'test to published TF relationships in an independent PBMC cohort.')

page_break()
h('Case study: the HCTL triad in hyperinflammatory diseases', 2)
para('We applied SCEPTIC to a candidate STAT4/NFKB1/GATA1 transcriptional triad in '
     'hyperinflammatory diseases. The triad was nominated from TRRUST v2²¹ TF enrichment '
     'analysis of our in-house discovery cohort and deliberately treated as an unvalidated '
     'cross-disease regulatory claim. The analytical question was therefore not whether '
     'the triad could be made significant in one dataset, but whether each of its three '
     'pairwise correlations (STAT4–NFKB1, NFKB1–GATA1, STAT4–GATA1) survived independent '
     'scoring, tissue, cohort, and null-model checks. Data sources and cohort '
     'characteristics are summarised in Table 2.')

# Table 2
t2_headers = ['Cohort', 'Disease', 'Assay', 'n', 'Role', 'Accession']
t2_rows = [
    ['Discovery (in-house)', 'MAS/Sepsis/IBD', 'PBMC scRNA-seq', '47', 'Discovery',
     'Available on request'],
    ['SCP548²⁶', 'Bacterial sepsis', 'PBMC scRNA-seq', '65', 'Validation (L1,L4,L5,L7)',
     'Broad SCP'],
    ['E-MTAB-10026²⁰', 'COVID-19', 'PBMC scRNA-seq', '143', 'Validation (L1,L7)',
     'E-MTAB-10026'],
    ['JIA (in-house)', 'Juvenile idiopathic arthritis', 'PBMC scRNA-seq', '4',
     'Exploratory (excluded from meta-analysis)', 'Available on request'],
    ['GSE65682 (MARS)²⁵', 'Sepsis (ICU)', 'Whole-blood microarray', '802',
     'Bulk replication (L3)', 'GSE65682'],
    ['Pienkos et al.²⁴', 'COVID-19', 'Plasma Olink + CyTOF', '100',
     'Cross-assay (L2)', 'OFID 2025; Zenodo 15199327'],
]
add_table(t2_headers, t2_rows,
          'Table 2. Cohort characteristics and data sources. Total unique patients: '
          '1,161 (6 cohorts, 3 assay types).',
          col_widths=[1.15, 1.0, 1.1, 0.35, 1.2, 1.0])

h('Layer 1: orthogonal scoring', 2)
para('Per-sample TF target-gene means were computed from TRRUST v2²¹ regulons (STAT4: '
     '9 targets; NFKB1: 17; GATA1: 10) for 47 discovery PBMC scRNA-seq samples '
     '(MAS n = 26, sepsis n = 12, IBD n = 9). SCENIC¹ AUCell served as the orthogonal '
     'method; DoRothEA confidence-weighted mean⁶, decoupleR multivariate linear model⁷, '
     'and GSVA²³ provide alternative scoring frameworks; PROGENy²² provides a '
     'pathway-level alternative. Raw target-gene means gave apparently significant or '
     'near-significant correlations for all three pairs: NFKB1–GATA1 r = +0.57 '
     '(p < 0.001), STAT4–NFKB1 r = −0.41 (p = 0.005), and STAT4–GATA1 r = −0.25 '
     '(p = 0.09). AUCell immediately separated method-stable from method-dependent '
     'signals: STAT4–NFKB1 retained the same direction and significance (r = −0.45, '
     'p = 0.002), whereas NFKB1–GATA1 collapsed to r = +0.15 (p = 0.33) and '
     'STAT4–GATA1 remained non-significant (r = −0.23, p = 0.18). Thus, Layer 1 '
     'rejected NFKB1–GATA1 and STAT4–GATA1 as discovery-level TF co-expression claims, '
     'leaving STAT4–NFKB1 as the only pair meriting cross-context testing (Figure 2).')

# Figure 2
insert_figure(os.path.join(OUTDIR, 'Fig2_final.png'), w=5.8,
              caption='Figure 2. NFKB1–GATA1 TF-score covariation: a method-dependent artifact under SCEPTIC Layer 1. Raw '
                      'target-gene mean scoring produces r = +0.57 (p < 0.001), but AUCell shows '
                      'r = +0.15 (p = 0.33). Housekeeping gene expression correlates with NFKB1 '
                      'targets at r = +0.97, consistent with global expression confounding.')

h('Caveat: GATA1 regulon specificity', 2)
para('The GATA1 regulon (TRRUST v2) includes erythroid-specific genes (HBB, HBA1, '
     'HBA2, ALAS2, GYPA, GYPB, KLF1). In PBMC and whole-blood samples, GATA1 '
     'target-gene scores primarily reflect erythrocyte contamination or ambient RNA '
     'rather than immune-cell GATA1 activity. Consequently, correlations involving '
     'GATA1 (NFKB1–GATA1, STAT4–GATA1) should be interpreted as reflecting shared '
     'erythroid signal rather than TF co-regulation in leukocytes. We recommend that '
     'SCEPTIC users inspect regulon composition for lineage-specific genes and report '
     'this as a potential confound.')

h('Layer 2: cross-assay replication', 2)
para('In an independent cohort (Pienkos et al.²⁴, n = 100), plasma Olink z-scores of '
     'NF-κB-related proteins (TNF, IL6, IL8, CXCL10) and STAT4/Th1-related proteins '
     '(IFNG, IL2) were positively correlated (r = +0.31, p = 0.002). CyTOF '
     'intracellular staining showed r = +0.80 (p < 1e-4) for the same pathway scores. '
     'These are cytokine-level readouts, not direct TF activity measurements; they '
     'should be interpreted only as downstream pathway support, not confirmatory '
     'evidence of TF nuclear activity, DNA binding, or regulon activity. No '
     'GATA1-related plasma proteins were available in the Olink panel, precluding '
     'Layer 2 testing of GATA1-involving pairs.')

h('Layer 3: bulk-tissue replication', 2)
para('In GSE65682 (MARS consortium²⁵, n = 802 whole-blood microarrays), the '
     'STAT4–NFKB1 correlation was +0.07 (p = 0.054), failing to replicate the '
     'scRNA-seq direction. NFKB1–GATA1 was −0.07 (p = 0.035), with opposite sign to '
     'the discovery raw-mean result. STAT4–GATA1 was −0.37 (p < 1e-4), but this '
     'likely reflects erythroid contamination in whole blood. All three pairs fail '
     'Layer 3.')

h('Layer 4: random TF null', 2)
para('We constructed the empirical null distribution of all N(N-1)/2 pairwise TF-score Pearson correlations among N = 32 TRRUST TFs expressed in the discovery cohort, yielding 496 background pairs. For the HCTL triad, we additionally computed the mean absolute pairwise |r| across the three TF pairs and compared it against 1,000 random TF triplets drawn from TRRUST v2²¹ (TFs with ≥ 5 targets in '
     'the gene universe) and computed their mean absolute pairwise correlation. The '
     'median mean |r| of random triplets was 0.457 (IQR 0.329–0.589, 95th percentile '
     '0.847). The observed HCTL triad mean |r| was 0.409, yielding an empirical '
     'p-value of 0.60—the HCTL triad is weaker than the majority of randomly chosen '
     'TFs. The mean |r| of 0.409 ranks at the 44th percentile of the triplet null '
     'distribution. As the broader literature audit demonstrates, L4 failure in mixed '
     'PBMC is common even for experimentally validated TF relationships; the key '
     'interpretation is not that the triad is uncorrelated, but that its correlation '
     'magnitude is unremarkable among TRRUST TFs. Sensitivity analysis using '
     'regulon-size-matched and expression-level-matched nulls yielded similar '
     'percentile ranks (mean Δ < 2%, Supplementary Figure S2), confirming that the '
     'random null is robust to TF set composition.')

h('Layer 5: composition-adjusted correlation (SCP548)', 2)
para('Linear regression of per-sample STAT4 and NFKB1 scores on coarse cell-type '
     'proportions (NK, T, B, myeloid, dendritic) in SCP548²⁶ (n = 65 donors) showed '
     'that NK cell and T cell proportions are the primary drivers of the per-sample '
     'correlation. The partial correlation after composition adjustment remained '
     'significant (r = 0.49, p < 1e-4), suggesting the covariation is not purely '
     'composition-driven. Reference-based deconvolution tools (CIBERSORTx²⁷, EPIC²⁸) '
     'could provide finer resolution, but their performance on scRNA-seq pseudobulk '
     'remains dataset-dependent²⁹. Residual confounding from unmeasured cell subtypes '
     'cannot be excluded.')

h('Cross-cohort replication of the HCTL triad', 2)
para('To assess the generalisability of the HCTL triad findings, we computed '
     'STAT4–NFKB1, NFKB1–GATA1, and STAT4–GATA1 per-sample correlations across all '
     'independent cohorts spanning three disease contexts (Figure 3A, Supplementary '
     'Figure S1). This analysis clarifies an apparent direction inconsistency. The '
     'pooled discovery STAT4–NFKB1 correlation was negative (r = −0.410), but '
     'per-disease decomposition showed positive within-disease correlations (MAS '
     'r = +0.69, IBD r = +0.93, sepsis r = +0.30), consistent with Simpson\'s paradox '
     'from pooling heterogeneous disease groups. In independent PBMC scRNA-seq cohorts, '
     'STAT4–NFKB1 was positive (SCP548 r = +0.561; E-MTAB-10026 r = +0.539), supporting '
     'a PBMC-context covariation signal. In GSE65682 whole-blood bulk microarray '
     '(n = 802), however, the correlation collapsed to r = +0.068 (p = 0.054), arguing '
     'against a general blood-wide regulatory axis. NFKB1–GATA1 attenuated across '
     'independent cohorts (Discovery r = +0.568; SCP548 r = +0.253; E-MTAB r = +0.049) '
     'and reversed in bulk (GSE65682 r = −0.074), consistent with method and erythroid '
     'signal confounding. STAT4–GATA1 was near zero in PBMC cohorts and negative in '
     'bulk. The cross-cohort pattern therefore narrows the case-study conclusion: one '
     'PBMC-context STAT4–NFKB1 signal remains, whereas the full HCTL triad does not.')

# Figure 3
insert_figure(os.path.join(OUTDIR, 'Fig3_final.png'), w=5.8,
              caption='Figure 3. Cross-cohort replication and cross-assay evidence. (A1–A3) Cross-cohort Pearson r for each HCTL TF pair across discovery and validation cohorts. STAT4–NFKB1 is consistent in PBMC only; NFKB1–GATA1 attenuates across cohorts; STAT4–GATA1 shows no reproducible signal. (B) STAT4–NFKB1 leave-one-cohort-out sensitivity (DerSimonian-Laird random-effects meta-analysis). (C) Olink and CyTOF protein-level pathway support for STAT4–NFKB1 (supportive evidence; not direct TF activity measurement).')

h('Layer 6: leave-one-cohort-out', 2)
para('Random-effects meta-analysis (DerSimonian-Laird, with Pearson r first transformed '
     'to Fisher z for variance stabilisation and back-transformed for reporting; '
     'metafor v4.1³⁹) across 6 cohort-modality combinations was used as a '
     'heterogeneity and dependence sensitivity analysis rather than as a direct '
     'common-effect biological estimate across assays. This analysis gave I² = 46% '
     '(Cochran Q = 119.1, df = 5, p < 1e-4), indicating substantial heterogeneity. '
     'Removing the two PBMC scRNA-seq cohorts (SCP548, E-MTAB-10026) rendered the '
     'pooled r non-significant. The positive finding is therefore cohort-dependent.')

h('Layer 7: label permutation', 2)
para('Layer 7 tests whether the observed within-cohort TF-score correlation is '
     'dependent on disease-group structure under an exchangeability null. Within '
     'each cohort, disease labels were randomly permuted 10,000 times while TF scores '
     'were held fixed, and the TF-pair Pearson r was recomputed for each permutation. '
     'The null hypothesis is that disease-group assignment does not contribute to the '
     'correlation magnitude. The observed STAT4–NFKB1 correlation exceeded the 95th '
     'percentile of the permutation null in SCP548 (permutation p < 0.001) and '
     'E-MTAB-10026 (p < 0.001), indicating that disease-status grouping contributes '
     'to the observed covariation. We note that L7 does not test whether the TF pair '
     'itself is specific, mechanistic, or causal; it only tests whether the observed '
     'covariation is compatible with exchangeable disease labels. A pair may pass L7 '
     'while failing the more stringent specificity test of L4.')

# Table 3: verdict
page_break()
t3_headers = ['TF pair', 'L1\nOrthogonal', 'L2\nCross-assay\n(supportive)', 'L3\nBulk',
              'L4\nRandom null', 'L5\nComposition', 'L6\nLOCO', 'L7\nPerm', 'Verdict']
t3_rows = [
    ['STAT4–NFKB1', '✓ Pass', '~ Supportive', '✗ Fail', '✗ Fail',
     '✓ Pass', '✗ Fail', '✓ Pass', 'CONTEXT-DEPENDENT\n(PBMC contexts only)'],
    ['NFKB1–GATA1', '✗ Fail', '— (N/A)', '✗ Fail', '✗ Fail',
     '✗ Fail', '✗ Fail', '—', 'REJECTED\n(method-dependent artifact)'],
    ['STAT4–GATA1', '✗ Fail', '— (N/A)', '✗ Fail', '✗ Fail',
     '✗ Fail', '✗ Fail', '—', 'REJECTED\n(inconsistent across methods)'],
]
add_table(t3_headers, t3_rows,
          'Table 3. SCEPTIC assessment for the HCTL triad. ✓ = layer passed; '
          '✗ = layer failed; — = not applicable or not tested.',
          col_widths=[0.85, 0.5, 0.5, 0.4, 0.5, 0.5, 0.4, 0.4, 1.1])


# ============================================================
# LITERATURE AUDIT
# ============================================================
page_break()
h('SCEPTIC validation of published TF co-expression claims', 2)

para('To assess whether SCEPTIC generalises beyond our own HCTL case study, we '
     'applied SCEPTIC layer 4 (random TF pair null) to a panel of 15 published TF '
     'co-expression claims from the peer-reviewed literature, spanning IFN signalling '
     '(STAT1–IRF1), T cell differentiation (IRF4–BATF, RUNX1–ETS1), myeloid '
     'development (SPI1–CEBPA), the NF-κB pathway (RELA–NFKB1), the AP-1 complex '
     '(FOS–JUN), Wnt signalling (TCF7–LEF1), the germinal centre B cell axis '
     '(BCL6–PRDM1), cell cycle regulation (MYC–E2F1), oxidative stress crosstalk '
     '(NFE2L2–HIF1A), Treg/Th2 balance (FOXP3–GATA3), glucocorticoid receptor–NF-κB '
     'crosstalk (NR3C1–RELA), and the TIL differentiation network (BATF–BCL6, '
     'Green et al. Immunity 2025). We computed per-sample TF target-gene mean scores '
     'for 32 TRRUST TFs in the SCP548 PBMC cohort (n = 65 donors, 496 TF pairs) and '
     'compared the absolute Pearson correlation of each literature pair against the '
     'empirical null distribution of all pairwise TF correlations (Figure 5, '
     'Supplementary Table S2).')

para('The null distribution showed a median |r| of 0.442 (IQR 0.293–0.616, '
     '95th percentile 0.847), confirming that most TF pairs exhibit moderate '
     'per-sample correlation in PBMC simply because many TF target-gene sets are '
     'driven by shared cell-type composition and global transcriptional activity. '
     'Of the 15 literature pairs tested, 5 passed SCEPTIC L4 at the top-5% threshold '
     '(STAT1–IRF1, RUNX1–ETS1, FOS–JUN, NFE2L2–HIF1A, NR3C1–RELA), 3 were marginal '
     '(IRF4–BATF, RELA–NFKB1, STAT1–IRF4), and 7 did not exceed the 75th percentile '
     'of all TF pairs (Figure 5).')

para('Importantly, the L4 results are biologically interpretable and do not invalidate '
     'the underlying TF relationships. BCL6–PRDM1 (|r| = 0.198, 21st percentile) are '
     'known transcriptional antagonists expressed in different B cell differentiation '
     'states; their low correlation in PBMC is expected. SPI1–CEBPA (|r| = 0.602, '
     '73rd percentile) are co-expressed in myeloid cells, but myeloid cells constitute '
     'only ~10% of PBMC, diluting their per-sample correlation below the top-5% '
     'threshold. SPI1–IRF8 (|r| = 0.284, 31st percentile) and FOXP3–GATA3 '
     '(|r| = 0.297, 32nd percentile) similarly reflect lineage-restricted expression '
     'patterns that do not produce pair-specific per-sample covariation in mixed '
     'PBMC. These results highlight a critical insight: even well-established, '
     'experimentally validated TF relationships do not necessarily produce per-sample '
     'co-expression signals that are distinguishable from background in heterogeneous '
     'tissue samples. SCEPTIC L4 provides a quantitative framework for making this '
     'distinction explicit.')

para('Several caveats apply to the literature audit. First, SCEPTIC L4 evaluates '
     'the specificity of per-sample TF score covariation in a given tissue context '
     '(here, PBMC), not the biological validity of the underlying TF relationship. '
     'A pair that does not exceed the 75th percentile in PBMC (e.g., BCL6–PRDM1, '
     'expressed in mutually exclusive B cell states) may be perfectly specific in '
     'purified cell populations. Second, the L4 null distribution depends on the '
     'regulon database (TRRUST v2), the expressed gene universe, and the sample '
     'cohort; results are regulon-and-context-conditional. Third, pairs exceeding '
     'the 95th percentile in PBMC may reflect shared cell-type co-expression rather '
     'than cell-intrinsic co-regulation. Despite these limitations, the audit '
     'demonstrates that SCEPTIC provides a standardised, quantitative framework for '
     'assessing whether a per-sample TF co-expression claim is distinguishable from '
     'background—a question rarely addressed in published single-cell studies.')

# Figure 5: literature audit
insert_figure(os.path.join(OUTDIR, 'Fig5_final.png'), w=6.0,
              caption='Figure 5. SCEPTIC L4 literature audit in SCP548 PBMC. (A) Empirical null '
                      'distribution of absolute TF-score correlations among 496 TRRUST TF pairs '
                      '(n = 65 donors). Dashed lines indicate median, 75th, and 95th percentiles. '
                      'Selected literature pairs overlaid. (B) Fifteen literature TF pairs ranked '
                      'by percentile in the PBMC null distribution. Green = top 5%, orange = '
                      '75th–95th percentile, red = below 75th percentile. L4 evaluates PBMC-context '
                      'TF-score covariation specificity only and does not invalidate the biological '
                      'TF relationship in other cell or tissue contexts.')

# ============================================================
# DISCUSSION
# ============================================================
page_break()
h('Discussion')

para('This study is primarily a methods and validation paper, not a claim that the '
     'HCTL triad is a universal hyperinflammatory-disease mechanism. We presented '
     'SCEPTIC, a seven-layer framework for validating single-cell TF co-expression '
     'claims, and used the HCTL triad as a stress test because it contains exactly the '
     'features that can mislead per-sample TF-score analyses: heterogeneous PBMC '
     'composition, disease severity structure, method-dependent scoring, and '
     'lineage-specific target genes. We then applied the same empirical-null principle '
     'to 15 published TF co-expression claims to show that the framework generalises '
     'beyond the case study.')

para('The HCTL case study yielded three lessons that define the practical logic of '
     'SCEPTIC. First, raw target-gene mean scoring '
     'can produce highly significant p-values for correlations that are not reproduced '
     'under regulon-based scoring or in bulk tissue (NFKB1–GATA1). This is consistent '
     'with known vulnerabilities of per-sample aggregate scoring¹¹⁻¹⁵ and benchmark '
     'studies showing method-dependent TF activity estimates³¹, and reinforces the '
     'need for orthogonal validation. Second, even when a correlation survives '
     'layers 1–2 in the discovery context, it may be weaker than randomly chosen TF '
     'tuples (STAT4–NFKB1 at layer 4, empirical p = 0.60) and substantially '
     'heterogeneous across cohorts (layer 6, I² = 46%). Third, the surviving signal '
     'is narrower than the original framing: not a universal cross-disease regulatory '
     'axis, but context-dependent covariation in PBMC inflammatory cohorts (sepsis, '
     'COVID-19, MAS, IBD).')

para('The random TF null (layer 4) is methodologically novel in the TF co-expression '
     'context but builds on established principles from permutation-based statistical '
     'inference³²,³³. In the single-cell field, statistical significance is typically '
     'assessed gene-by-gene, and a significant p-value for a TF pair correlation is '
     'taken as evidence of biological coordination. The layer 4 null asks a different '
     'question: is this specific pair more correlated than randomly chosen TFs with '
     'comparable target-gene support? For the HCTL triad, the answer was no—the '
     'observed mean |r| of 0.409 fell below the null median of 0.457. This does not '
     'mean STAT4 and NFKB1 are uncorrelated; it means their correlation magnitude is '
     'not unusual among TRRUST TFs in these samples, and cannot be used to claim a '
     'special regulatory relationship. The empirical null distribution depends on '
     'regulon quality²¹ and the expressed gene universe; results should be interpreted '
     'as regulon-database-conditional.')

para('The literature audit separates the framework from the HCTL example. Of 15 published TF pairs '
     'tested, only 5 exceeded the 95th percentile of the null distribution in SCP548 '
     'PBMC. Pairs that did not exceed the 75th percentile included several with '
     'extensive independent experimental validation (SPI1–CEBPA, BCL6–PRDM1, '
     'FOXP3–GATA3)—not because these relationships are false, but because their '
     'per-sample covariation in mixed PBMC is not pair-specific. This distinction is '
     'precisely what SCEPTIC L4 is designed to make. Authors reporting TF co-expression '
     'from per-sample scores should routinely apply L4 to contextualise their findings '
     'within the broader landscape of TF co-variation.')

para('We recommend that single-cell studies reporting per-sample TF co-expression '
     'programs adopt SCEPTIC or an equivalent multi-layer validation framework. At '
     'minimum, authors should report layer 1 (orthogonal scoring) and layer 4 '
     '(random TF null), because these two checks directly address the most common '
     'failure mode: a nominally significant TF-score correlation that is neither '
     'method-stable nor pair-specific. For layers requiring independent datasets '
     '(L2, L3), absence should be explicitly noted as a limitation. We emphasise that '
     'SCEPTIC verdicts are specific to the tissue context, regulon database, and '
     'scoring method used; a pair passing all seven layers may still be a correlative '
     'bystander rather than a causal driver.')

para('Several limitations should be noted. The cytokine-level assays in Layer 2 '
     'measure downstream effectors rather than TF activity directly. The composition '
     'adjustment in Layer 5 uses coarse cell-type proportions from scRNA-seq '
     'annotations; residual confounding from unmeasured subtypes is possible. The '
     'framework is designed for per-sample aggregate TF scoring; it does not address '
     'single-cell-level regulatory network inference methods. The random null '
     '(Layer 4) depends on the completeness of the chosen regulon database and can be '
     'affected by target overlap or lineage-restricted target genes. Finally, '
     'SCEPTIC assesses the validity of co-expression claims, not their biological '
     'importance.')

para('In summary, SCEPTIC converts TF co-expression interpretation from a single '
     'correlation-and-p-value claim into a structured validation decision. The framework '
     'complements existing best-practice guidelines¹⁴,¹⁵ and '
     'differential expression benchmarking studies³⁴,³⁵ by specifically addressing '
     'the validation of TF co-expression claims—a common but under-scrutinised '
     'analysis step. The checklist is provided as an open-source resource '
     '(Supplementary Table S1). We encourage its adoption, extension, and independent '
     'evaluation.')

# ============================================================
# METHODS
# ============================================================
page_break()
h('Methods')

h('Data sources and preprocessing', 2)
para('Six retrospective cohorts comprising 1,161 unique patients were analysed '
     '(Table 2). For each scRNA-seq dataset, raw count matrices were downloaded from '
     'the respective repository. Cells were filtered to retain those with 200–6,000 '
     'detected genes, fewer than 20% mitochondrial reads, and more than 500 total '
     'UMI counts. Doublet removal was performed per dataset using Scrublet (default '
     'parameters). Counts were library-size normalised to 10,000 counts per cell and '
     'log1p-transformed (scanpy.pp.normalize_total, scanpy.pp.log1p). Per-sample '
     'pseudobulk scores were computed as the arithmetic mean of log-normalised '
     'expression across all cells from the same donor. Batch correction was not '
     'applied to per-sample aggregate scores, as each cohort was analysed '
     'independently; cross-cohort comparisons used cohort-specific null distributions. '
     'Each analytical row represents one donor/sample-level summary; where source '
     'studies contained repeated sampling designs, repeated-measures inference was not '
     'claimed unless repeated samples were explicitly modelled by the source study. '
     'Complete per-cohort QC metrics are provided in Supplementary Table S3.')

para('MAS/Sepsis/IBD (n = 47): in-house PBMC scRNA-seq (Nanfang Hospital IRB '
     'approved). SCP548 (n = 65): Broad Single Cell Portal²⁶, bacterial sepsis '
     'PBMC scRNA-seq (10x Genomics platform³⁶). E-MTAB-10026 (n = 143): COVID-19 '
     'PBMC scRNA-seq, EBI ArrayExpress²⁰. JIA (n = 4): in-house PBMC scRNA-seq, '
     'exploratory only, excluded from quantitative meta-analysis. GSE65682 (n = 802): '
     'MARS consortium whole-blood microarrays, NCBI GEO²⁵. Pienkos et al. (n = 100): '
     'plasma Olink + CyTOF, Open Forum Infectious Diseases / Zenodo²⁴. Public datasets were de-identified at source '
     'and used in accordance with data access policies.')

h('TF pathway scoring', 2)
para('TRRUST v2²¹ regulons were used: STAT4 (9 targets: IFNG, IL2, IL12RB2, TBX21, '
     'CCR5, CXCR3, IL18R1, IL18RAP, HAVCR2), NFKB1 (17 targets: TNF, IL6, CXCL8, '
     'IL1B, CCL2, CCL3, CCL4, CCL5, CXCL10, ICAM1, VCAM1, SELE, PTGS2, MMP9, BCL2, '
     'BCL2L1, SOD2), GATA1 (10 targets: HBB, HBA1, HBA2, ALAS2, EPB42, GYPA, GYPB, '
     'KLF1, NFE2, SP1). Per-sample mean expression of target genes was computed from '
     'log-normalised count matrices. SCENIC AUCell scores were computed via '
     'pySCENIC⁵ for the discovery cohort as an orthogonal method (Layer 1). Plasma '
     'Olink NPX values were normalised using reference-sample bridging³⁷; CyTOF MFI '
     'values were arcsinh-transformed³⁸. Pathway scores were computed as the mean of '
     'available target proteins within each functional category. Target overlap, '
     'lineage specificity, missing target genes, and regulon database quality can '
     'alter TF-score correlations; SCEPTIC therefore treats L4 results as '
     'regulon-and-context-conditional rather than as direct tests of TF biology.')

h('SCEPTIC layer implementation', 2)
para('Layer 1: Pearson correlation under both raw target-gene mean and SCENIC AUCell. '
     'A pair passes if both methods yield the same sign, p < 0.05, and |r| > 0.3. '
     'Layer 2: Pearson correlation of Olink/CyTOF pathway z-scores. Layer 3: Pearson '
     'correlation in GSE65682 bulk microarray. Layer 4: the observed TF-pair |r| is '
     'compared against the empirical distribution of all N(N-1)/2 pairwise TF-score '
     'correlations among expressed TRRUST TFs (pair-level null). Because L4 uses '
     'absolute r, positive and negative score covariation are combined for the '
     'specificity screen; biological interpretation therefore returns to signed r and '
     'the known TF relationship. For the HCTL triad-level '
     'summary, we additionally compare the mean pairwise |r| against 1,000 random TF '
     'triplets drawn from TRRUST v2 (TFs with ≥ 5 targets in the expressed gene universe); '
     'empirical p = fraction with mean absolute pairwise r ≥ observed. Layer 5: '
     'partial correlation after linear regression of TF scores on coarse cell-type '
     'proportions (NK, T, B, myeloid, dendritic) and housekeeping gene mean '
     'expression. Layer 6: random-effects meta-analysis, with Pearson r transformed '
     'to Fisher z for variance stabilisation and back-transformed for reporting '
     '(metafor v4.1³⁹), across 6 cohort-modality combinations as a descriptive '
     'heterogeneity check; leave-one-cohort-out '
     'sensitivity; Cochran Q and I² for heterogeneity. Layer 7: disease labels '
     'permuted 10,000 times within each cohort; null hypothesis is that TF-pair '
     'correlation does not depend on disease assignment; observed |r| compared '
     'against permutation null. Pass criteria for each layer are given in Table 1 and '
     'are intended as reproducible reporting heuristics, not universal biological '
     'truths.')

h('Matched null sensitivity analysis', 2)
para('To assess the robustness of SCEPTIC L4 classifications to null design, we '
     'compared the global (all-TF) null distribution against two matched nulls: '
     '(1) regulon-size-matched: TFs with ± 3 target genes of the test pair, and '
     '(2) expression-level-matched: TFs with mean expression within ± 0.5 SD of '
     'the test pair. For each test pair, 500 random pairs were sampled from each '
     'matched pool; resulting percentile ranks were compared against the global null. '
     'The mean difference between global and size-matched percentile was +1.6% '
     '(n = 16 pairs), and between global and expression-matched was −0.0%, indicating '
     'that L4 classification is robust to null specification (Supplementary Figure S2, '
     'Supplementary Table S4).')

h('Literature audit', 2)
para('For the literature audit, 15 published TF pairs were selected to span diverse '
     'biological contexts and varying levels of experimental validation. Per-sample '
     'TF target-gene mean scores were computed for 32 TRRUST v2 TFs (≥ 3 targets '
     'expressed) in the SCP548 PBMC cohort (n = 65 donors). All 496 pairwise Pearson '
     'correlations were computed to form the empirical null distribution. Each '
     'literature pair was compared against this null; pairs with |r| exceeding the '
     '95th percentile were classified as SCEPTIC L4 PASS, those exceeding the 75th '
     'percentile as MARGINAL, and the remainder as not exceeding the 75th percentile '
     '(see Figure 5). Full audit results are in Supplementary Table S2.')

h('Statistical analysis', 2)
para('Pearson correlation was used throughout; Spearman rank correlation was computed '
     'as a robustness check for all primary TF-pair correlations and yielded consistent '
     'results (Supplementary Table S5). P-values for individual correlations are '
     'two-sided. 95% confidence intervals for Pearson r were computed via Fisher z '
     'transformation and are reported for primary correlations in Supplementary Table '
     'S5. Multiple testing correction (Benjamini-Hochberg¹⁸) was applied within each '
     'layer where multiple TF pairs were tested (i.e., 3 TF pairs for the HCTL case '
     'study in L1, L3; 15 literature pairs in the L4 audit; 496 background pairs for '
     'the null distribution). A dedicated testing-family table is provided as '
     'Supplementary Table S6. Meta-analysis used the DerSimonian-Laird '
     'random-effects model³⁰. Permutation p-values (Layers 4 and 7) were computed as '
     '(k + 1)/(N + 1) where k = number of null replicates exceeding the observed '
     'statistic and N = number of permutations (thus minimum p = 1/1001 ≈ 0.001 for '
     'L4, and 1/10001 ≈ 1e-4 for L7). Bootstrap confidence intervals for I² used '
     '1,000 resamples³³. Analyses were performed in Python 3.12 (scanpy⁴⁰, scipy, '
     'statsmodels⁴¹, scikit-learn⁴²) and R 4.5 (metafor³⁹, sva⁴³). Dimensionality '
     'reduction for exploratory visualisation used UMAP⁴⁴. Complete random seeds are '
     'recorded in the analysis scripts.')

# ============================================================
# DATA AND CODE AVAILABILITY
# ============================================================
h('Data and code availability')
para('Public datasets: SCP548 (https://singlecell.broadinstitute.org/single_cell/study/SCP548), '
     'E-MTAB-10026 (https://www.ebi.ac.uk/biostudies/arrayexpress/studies/E-MTAB-10026), '
     'GSE65682 (NCBI GEO accession GSE65682), Pienkos et al. '
     '(https://doi.org/10.1093/ofid/ofaf515; data DOI: https://doi.org/10.5281/zenodo.15199327)²⁴. In-house PBMC scRNA-seq data are '
     'available from the corresponding author upon reasonable request and with '
     'institutional data access approval. Per-sample TF scores, literature audit '
     'results, and all processed intermediate data required to reproduce Figures 1–5 '
     'and Supplementary Figures S1–S2 are provided as Supplementary Data files '
     'accompanying this manuscript.')

para('Analysis scripts are available at https://github.com/xutaoguo55/SCEPTIC '
     '(private repository; access will be granted to editors and reviewers upon '
     'request, and the repository will be made public upon acceptance). A frozen '
     'version will be archived on Zenodo with a persistent DOI upon publication.')

# ============================================================
# SUPPLEMENTARY INFORMATION
# ============================================================
page_break()
h('Supplementary Information')

h('Supplementary Table S1: SCEPTIC checklist', 2)
para('The following checklist can be used by authors and reviewers to assess whether '
     'a reported TF co-expression claim has been adequately validated.')

s1_headers = ['#', 'Checklist item', 'Response options', 'Minimum standard']
s1_rows = [
    ['1', 'Is TF activity scored by ≥ 2 independent methods?',
     'Yes / No / N/A', 'Required reporting check (Layer 1)'],
    ['2', 'Are regulon sources, target gene lists, and gene universe fully reported?',
     'Yes / No', 'Required reporting check'],
    ['3', 'Is correlation direction consistent between scoring methods?',
     'Yes (same sign) / No', 'Required reporting check (Layer 1)'],
    ['4', 'For cross-disease claims: replication in ≥ 2 independent cohorts?',
     'Yes (≥ 2) / Partial (1) / No', 'Recommended (Layers 2, 6)'],
    ['5', 'Is protein-level support provided?',
     'Yes / No (noted as limitation)', 'Recommended (Layer 2)'],
    ['6', 'Is finding replicated in bulk tissue from independent cohort?',
     'Yes / No / N/A', 'Recommended (Layer 3)'],
    ['7', 'Is observed TF-pair |r| compared against random TF-pair null?',
     'Yes, with empirical p / No', 'Required reporting check (Layer 4)'],
    ['8', 'Are cell-type proportions regressed out?',
     'Yes, with partial r / No', 'Recommended (Layer 5)'],
    ['9', 'Is leave-one-cohort-out performed when ≥ 3 cohorts available?',
     'Yes / No / N/A', 'Recommended (Layer 6)'],
    ['10', 'Are disease-label permutation p-values reported?',
     'Yes / No', 'Recommended (Layer 7)'],
    ['11', 'If any layer is omitted, is omission explicitly noted as limitation?',
     'Yes / No', 'Required reporting check'],
    ['12', 'Are per-sample scores, metadata, and analysis scripts publicly available?',
     'Yes / Upon request / No', 'Required reporting check'],
]
add_table(s1_headers, s1_rows,
          'Supplementary Table S1. SCEPTIC validation checklist for reporting TF '
          'co-expression claims.',
          col_widths=[0.3, 3.0, 1.2, 1.2])

h('Supplementary Table S2: Literature audit results', 2)
para('SCEPTIC L4 applied to 15 published TF co-expression claims in SCP548 PBMC '
     '(n = 65 donors, 32 TFs, 496 TF pairs). Null: median |r| = 0.442, '
     '75th = 0.616, 95th = 0.847.')

s2_headers = ['TF pair', 'r', 'p', '|r|', '%ile', 'Emp p', 'L4 result', 'Source']
s2_rows = [
    ['STAT1–IRF1', '+0.947', '<1e-4', '0.947', '99.0', '0.010', 'PASS (top 5%)', 'JAK-STAT/IFN; Darnell Science 1997'],
    ['RUNX1–ETS1', '+0.967', '<1e-4', '0.967', '99.8', '0.002', 'PASS (top 5%)', 'T cell dev; Taniuchi Cell 2002'],
    ['NR3C1–RELA', '+0.926', '<1e-4', '0.926', '98.2', '0.018', 'PASS (top 5%)', 'GR-NFkB; Caldenhoven Mol Endo 1995'],
    ['FOS–JUN', '+0.882', '<1e-4', '0.882', '97.4', '0.026', 'PASS (top 5%)', 'AP-1; Angel BBA 1991'],
    ['NFE2L2–HIF1A', '+0.881', '<1e-4', '0.881', '97.0', '0.030', 'PASS (top 5%)', 'Stress crosstalk; Hayes ARPT 2020'],
    ['RELA–NFKB1', '+0.698', '<1e-4', '0.698', '84.1', '0.159', 'MARGINAL', 'NF-kB; Hayden Cell 2008'],
    ['IRF4–BATF', '+0.677', '<1e-4', '0.677', '80.6', '0.194', 'MARGINAL', 'Tfh diff.; Murphy Nat Rev Immunol 2013'],
    ['STAT1–IRF4', '+0.641', '<1e-4', '0.641', '78.0', '0.220', 'MARGINAL', 'Th1/Tfh balance'],
    ['STAT4–NFKB1', '+0.610', '<1e-4', '0.610', '74.2', '0.258', '< 75th %ile', 'HCTL triad; this study'],
    ['TCF7–LEF1', '+0.606', '<1e-4', '0.606', '73.2', '0.268', '< 75th %ile', 'Wnt/T cell; Staal Nat Rev Immunol 2018'],
    ['MYC–E2F1', '+0.603', '<1e-4', '0.603', '73.0', '0.270', '< 75th %ile', 'Cell cycle; Leone Nature 1997'],
    ['SPI1–CEBPA', '+0.602', '<1e-4', '0.602', '72.6', '0.274', '< 75th %ile', 'Myeloid core; Friedman Blood 2002'],
    ['BATF–BCL6', '+0.398', '0.001', '0.398', '45.0', '0.550', '< 75th %ile', 'TIL network; Green Immunity 2025⁴⁵'],
    ['FOXP3–GATA3', '+0.297', '0.016', '0.297', '32.3', '0.677', '< 75th %ile', 'Treg/Th2; Rudra Nature 2012'],
    ['SPI1–IRF8', '+0.284', '0.022', '0.284', '31.2', '0.688', '< 75th %ile', 'Myeloid dev; Tamura Annu Rev Immunol 2008'],
    ['BCL6–PRDM1', '+0.198', '0.114', '0.198', '20.8', '0.792', '< 75th %ile', 'GC B cell; Crotty Annu Rev Immunol 2010'],
]
add_table(s2_headers, s2_rows,
          'Supplementary Table S2. SCEPTIC L4 literature audit results (SCP548 PBMC).',
          col_widths=[0.65, 0.35, 0.3, 0.3, 0.3, 0.35, 0.6, 1.5])

h('Supplementary Table S3: Cohort QC and preprocessing', 2)
add_csv_table(os.path.join(BASE, 'results/literature_audit/table_s3_qc_metrics.csv'),
              'Supplementary Table S3. Cohort QC metrics, platform details, sample counts, '
              'and preprocessing parameters.',
              col_widths=[0.75, 0.75, 0.85, 0.85, 0.45, 0.55, 0.85, 1.05, 0.75, 0.55])

h('Supplementary Table S4: L4 null design sensitivity', 2)
add_csv_table(os.path.join(BASE, 'results/literature_audit/matched_null_sensitivity.csv'),
              'Supplementary Table S4. L4 null design sensitivity. Global percentile ranks '
              'are compared against regulon-size-matched and expression-level-matched nulls.',
              col_widths=[0.8, 0.3, 0.3, 0.45, 0.5, 0.5, 0.5, 0.45, 0.45, 0.35, 0.35])

h('Supplementary Table S5: Spearman robustness check', 2)
add_csv_table(os.path.join(BASE, 'results/literature_audit/table_s5_spearman_robustness.csv'),
              'Supplementary Table S5. Spearman robustness check for primary HCTL TF-pair '
              'correlations across scoring methods and cohorts, with Fisher-z 95% CI '
              'for Pearson r.',
              col_widths=[0.65, 0.6, 0.25, 0.38, 0.38, 0.38, 0.38, 0.38, 0.38, 0.45])

h('Supplementary Table S6: Statistical testing families', 2)
add_csv_table(os.path.join(BASE, 'SCEPTIC_Supplementary_Data/supplementary_tables/table_s6_statistical_testing_families.csv'),
              'Supplementary Table S6. Statistical testing families, effect sizes, p-value '
              'types, confidence-interval reporting, and interpretation limits for each '
              'SCEPTIC layer.',
              col_widths=[0.75, 0.75, 1.15, 0.85, 1.05, 0.85, 1.25])

h('Supplementary Table S7: Parameters and thresholds', 2)
add_csv_table(os.path.join(BASE, 'SCEPTIC_Supplementary_Data/supplementary_tables/table_s7_sceptic_parameters_thresholds.csv'),
              'Supplementary Table S7. SCEPTIC parameters and thresholds. Thresholds '
              'are intended as standardised reporting heuristics, not universal '
              'biological cutoffs.',
              col_widths=[0.5, 1.05, 2.45, 2.0])

h('Supplementary Table S8: Machine-readable checklist', 2)
add_csv_table(os.path.join(BASE, 'SCEPTIC_Supplementary_Data/supplementary_tables/machine_readable_sceptic_checklist.csv'),
              'Supplementary Table S8. Machine-readable SCEPTIC checklist. The same '
              'content is provided as CSV and JSON in the Supplementary Data archive '
              'and code repository resources directory.',
              col_widths=[0.35, 3.1, 0.65, 1.6])

# ============================================================
# DECLARATIONS
# ============================================================
page_break()
h('Declarations')

h('Acknowledgements', 2)
para('We thank the patients and their families for contributing samples to the '
     'original studies whose public data made this analysis possible. We thank the '
     'Broad Institute Single Cell Portal, ArrayExpress, and NCBI GEO for maintaining '
     'open data repositories. Computational resources were provided by the Nanfang '
     'Hospital Clinical Research Center.')

h('Funding', 2)
para('This work was supported by institutional funds from Nanfang Hospital, Southern '
     'Medical University. No external grant funding was received for this specific '
     'study.')

h('Competing interests', 2)
para('The authors declare no competing interests.')

h('Author contributions', 2)
para('X.G. conceived the study, developed the SCEPTIC framework, performed all '
     'computational analyses, and wrote the manuscript. X.W., H.Z., and H.J. '
     'contributed equally to data curation, cohort identification, and result '
     'interpretation. J.H., Q.W., and Y.W. contributed to data collection and '
     'validation. R.F. provided resources and supervision. All authors reviewed '
     'and approved the final manuscript.')

h('Data availability', 2)
para('All public datasets used in this study are available from their respective '
     'repositories (SCP548, E-MTAB-10026, GSE65682, OFID/Pienkos data DOI 10.5281/zenodo.15199327). Per-sample '
     'TF scores, literature audit results, and processed intermediate data are '
     'provided as Supplementary Data files accompanying this manuscript. In-house '
     'PBMC scRNA-seq data are available from the corresponding author upon reasonable '
     'request, subject to institutional data access approval.')

h('Code availability', 2)
para('Analysis scripts are available at https://github.com/xutaoguo55/SCEPTIC '
     '(private repository; access will be granted to editors and reviewers upon '
     'request, and the repository will be made public upon acceptance). A frozen '
     'version will be archived on Zenodo with a persistent DOI upon publication.')

# ============================================================
# REFERENCES
# ============================================================
page_break()
h('References')
REFS = [
    '1. Aibar, S. et al. SCENIC: single-cell regulatory network inference and clustering. Nat. Methods 14, 1083–1086 (2017). https://doi.org/10.1038/nmeth.4463.',
    '2. Lopez, R., Regier, J., Cole, M.B., Jordan, M.I. & Yosef, N. Deep generative modeling for single-cell transcriptomics. Nat. Methods 15, 1053–1058 (2018). https://doi.org/10.1038/s41592-018-0229-2.',
    '3. Stuart, T. et al. Comprehensive integration of single-cell data. Cell 177, 1888–1902.e21 (2019). https://doi.org/10.1016/j.cell.2019.05.031.',
    '4. Hao, Y. et al. Dictionary learning for integrative, multimodal and scalable single-cell analysis. Nat. Biotechnol. 42, 293–304 (2024). https://doi.org/10.1038/s41587-023-01767-y.',
    '5. Van de Sande, B. et al. A scalable SCENIC workflow for single-cell gene regulatory network analysis. Nat. Protoc. 15, 2247–2276 (2020). https://doi.org/10.1038/s41596-020-0336-2.',
    '6. Garcia-Alonso, L. et al. Benchmark and integration of resources for the estimation of human transcription factor activities. Genome Res. 29, 1363–1375 (2019). https://doi.org/10.1101/gr.240663.118.',
    '7. Badia-i-Mompel, P. et al. decoupleR: ensemble of computational methods to infer biological activities from omics data. Bioinform. Adv. 2, vbac016 (2022). https://doi.org/10.1093/bioadv/vbac016.',
    '8. Alvarez, M.J. et al. Functional characterization of somatic mutations in cancer using network-based inference of protein activity. Nat. Genet. 48, 838–847 (2016). https://doi.org/10.1038/ng.3593.',
    '9. Subramanian, A. et al. Gene set enrichment analysis: a knowledge-based approach for interpreting genome-wide expression profiles. Proc. Natl Acad. Sci. USA 102, 15545–15550 (2005). https://doi.org/10.1073/pnas.0506580102.',
    '10. Barbie, D.A. et al. Systematic RNA interference reveals that oncogenic KRAS-driven cancers require TBK1. Nature 462, 108–112 (2009). https://doi.org/10.1038/nature08460.',
    '11. Squair, J.W. et al. Confronting false discoveries in single-cell differential expression. Nat. Commun. 12, 5692 (2021). https://doi.org/10.1038/s41467-021-25960-2.',
    '12. Finak, G. et al. MAST: a flexible statistical framework for assessing transcriptional changes and characterizing heterogeneity in single-cell RNA sequencing data. Genome Biol. 16, 278 (2015). https://doi.org/10.1186/s13059-015-0844-5.',
    '13. Kharchenko, P.V., Silberstein, L. & Scadden, D.T. Bayesian approach to single-cell differential expression analysis. Nat. Methods 11, 740–742 (2014). https://doi.org/10.1038/nmeth.2967.',
    '14. Luecken, M.D. & Theis, F.J. Current best practices in single-cell RNA-seq analysis: a tutorial. Mol. Syst. Biol. 15, e8746 (2019). https://doi.org/10.15252/msb.20188746.',
    '15. Heumos, L. et al. Best practices for single-cell analysis across modalities. Nat. Rev. Genet. 24, 550–572 (2023). https://doi.org/10.1038/s41576-023-00586-w.',
    '16. Lähnemann, D. et al. Eleven grand challenges in single-cell data science. Genome Biol. 21, 31 (2020). https://doi.org/10.1186/s13059-020-1926-6.',
    '17. Hotchkiss, R.S. et al. Sepsis and septic shock. Nat. Rev. Dis. Primers 2, 16045 (2016). https://doi.org/10.1038/nrdp.2016.45.',
    '18. Singer, M. et al. The Third International Consensus Definitions for Sepsis and Septic Shock (Sepsis-3). JAMA 315, 801–810 (2016). https://doi.org/10.1001/jama.2016.0287.',
    '19. Venet, F. & Monneret, G. Advances in the understanding and treatment of sepsis-induced immunosuppression. Nat. Rev. Nephrol. 14, 121–137 (2018). https://doi.org/10.1038/nrneph.2017.165.',
    '20. Stephenson, E. et al. Single-cell multi-omics analysis of the immune response in COVID-19. Nat. Med. 27, 904–916 (2021). https://doi.org/10.1038/s41591-021-01329-2.',
    '21. Han, H. et al. TRRUST v2: an expanded reference database of human and mouse transcriptional regulatory interactions. Nucleic Acids Res. 46, D380–D386 (2018). https://doi.org/10.1093/nar/gkx1013.',
    '22. Schubert, M. et al. Perturbation-response genes reveal signaling footprints in cancer gene expression. Nat. Commun. 9, 20 (2018). https://doi.org/10.1038/s41467-017-02391-6.',
    '23. Hänzelmann, S., Castelo, R. & Guinney, J. GSVA: gene set variation analysis for microarray and RNA-seq data. BMC Bioinformatics 14, 7 (2013). https://doi.org/10.1186/1471-2105-14-7.',
    '24. Pienkos, S. et al. Single-cell and plasma proteomics do not differentiate patients with and without SARS-CoV-2 antigenemia in convalescence in a cohort of 100 patients. Open Forum Infect. Dis. 12, ofaf515 (2025). https://doi.org/10.1093/ofid/ofaf515; data DOI: https://doi.org/10.5281/zenodo.15199327.',
    '25. Scicluna, B.P. et al. Classification of patients with sepsis according to blood genomic endotype: a prospective cohort study. Lancet Respir. Med. 5, 816–826 (2017). https://doi.org/10.1016/s2213-2600(17)30294-1.',
    '26. Reyes, M. et al. An immune-cell signature of bacterial sepsis. Nat. Med. 26, 333–340 (2020). https://doi.org/10.1038/s41591-020-0752-4.',
    '27. Newman, A.M. et al. Determining cell type abundance and expression from bulk tissues with digital cytometry. Nat. Biotechnol. 37, 773–782 (2019). https://doi.org/10.1038/s41587-019-0114-2.',
    '28. Racle, J. & Gfeller, D. EPIC: a tool to estimate the proportions of different cell types from bulk gene expression data. Methods Mol. Biol. 2120, 233–248 (2020). https://doi.org/10.1007/978-1-0716-0327-7_17.',
    '29. Avila Cobos, F., Alquicira-Hernandez, J., Powell, J.E., Mestdagh, P. & De Preter, K. Benchmarking of cell type deconvolution pipelines for transcriptomics data. Nat. Commun. 11, 5650 (2020). https://doi.org/10.1038/s41467-020-19015-1.',
    '30. DerSimonian, R. & Laird, N. Meta-analysis in clinical trials. Control. Clin. Trials 7, 177–188 (1986). https://doi.org/10.1016/0197-2456(86)90046-2.',
    '31. Holland, C.H., Tanevski, J., Perales-Patón, J. et al. Robustness and applicability of transcription factor and pathway analysis tools on single-cell RNA-seq data. Genome Biol. 21, 36 (2020). https://doi.org/10.1186/s13059-020-1949-z.',
    '32. Fisher, R.A. The Design of Experiments. (Oliver & Boyd, Edinburgh, 1935).',
    '33. Efron, B. & Tibshirani, R.J. An Introduction to the Bootstrap. (Chapman & Hall, New York, 1993).',
    '34. Soneson, C. & Robinson, M.D. Bias, robustness and scalability in single-cell differential expression analysis. Nat. Methods 15, 255–261 (2018). https://doi.org/10.1038/nmeth.4612.',
    '35. Love, M.I., Huber, W. & Anders, S. Moderated estimation of fold change and dispersion for RNA-seq data with DESeq2. Genome Biol. 15, 550 (2014). https://doi.org/10.1186/s13059-014-0550-8.',
    '36. Zheng, G.X.Y. et al. Massively parallel digital transcriptional profiling of single cells. Nat. Commun. 8, 14049 (2017). https://doi.org/10.1038/ncomms14049.',
    '37. Assarsson, E. et al. Homogenous 96-plex PEA immunoassay exhibiting high sensitivity, specificity, and excellent scalability. PLoS ONE 9, e95192 (2014). https://doi.org/10.1371/journal.pone.0095192.',
    '38. Bendall, S.C. et al. Single-cell mass cytometry of differential immune and drug responses across a human hematopoietic continuum. Science 332, 687–696 (2011). https://doi.org/10.1126/science.1198704.',
    '39. Viechtbauer, W. Conducting meta-analyses in R with the metafor package. J. Stat. Softw. 36, 1–48 (2010). https://doi.org/10.18637/jss.v036.i03.',
    '40. Wolf, F.A., Angerer, P. & Theis, F.J. SCANPY: large-scale single-cell gene expression data analysis. Genome Biol. 19, 15 (2018). https://doi.org/10.1186/s13059-017-1382-0.',
    '41. Seabold, S. & Perktold, J. Statsmodels: econometric and statistical modeling with Python. Proc. 9th Python Sci. Conf. 92–96 (2010). https://doi.org/10.25080/Majora-92bf1922-011.',
    '42. Pedregosa, F. et al. Scikit-learn: machine learning in Python. J. Mach. Learn. Res. 12, 2825–2830 (2011).',
    '43. Leek, J.T. & Storey, J.D. Capturing heterogeneity in gene expression studies by surrogate variable analysis. PLoS Genet. 3, e161 (2007). https://doi.org/10.1371/journal.pgen.0030161.',
    '44. McInnes, L., Healy, J. & Melville, J. UMAP: Uniform Manifold Approximation and Projection for Dimension Reduction. arXiv:1802.03426 (2018).',
    '45. Green, W.D. et al. Enhancer-driven gene regulatory networks reveal transcription factors governing T cell adaptation and differentiation in the tumor microenvironment. Immunity 58, 1725–1741.e9 (2025). https://doi.org/10.1016/j.immuni.2025.04.030.',
]
for ref in REFS:
    pp = doc.add_paragraph()
    pp.paragraph_format.space_after = Pt(1)
    rr = pp.add_run(ref)
    rr.font.size = Pt(9)

# ========== SAVE ==========
out = os.path.join(BASE, 'SCEPTIC_SUBMISSION.docx')
doc.save(out)
print(f"Saved: {out} ({len(REFS)} refs)")
